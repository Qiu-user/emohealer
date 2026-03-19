#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成 EmoHealer HLD 文档
基于 EmoHealer_HLD_v1.0.md 生成 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建文档
doc = Document()

# 设置页面样式
section = doc.sections[0]
section.page_height = Cm(29.7)  # A4高度
section.page_width = Cm(21.0)   # A4宽度
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# 设置中文字体
def set_chinese_font(run, font_name='微软雅黑', size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

# 添加标题
title = doc.add_heading('EmoHealer 情绪疗愈平台', level=0)
run = title.runs[0]
set_chinese_font(run, '微软雅黑', 20, True)

subtitle = doc.add_heading('高层设计文档 (HLD)', level=1)
run = subtitle.runs[0]
set_chinese_font(run, '微软雅黑', 16, False)

# 添加文档信息
info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Table Grid'

info_data = [
    ['项目名称', 'EmoHealer AI 情绪疗愈平台'],
    ['文档版本', 'v1.0'],
    ['创建日期', '2026-03-19'],
    ['作者', '基于 EmoHealer 项目代码分析生成'],
    ['对应 SRS', 'EmoHealer_功能需求文档.md'],
    ['文档类型', 'High-Level Design (HLD)']
]

for i, (key, value) in enumerate(info_data):
    row = info_table.rows[i]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(key), '微软雅黑', 11, True)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(value), '微软雅黑', 11, False)
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(4.0)

doc.add_paragraph()

# 添加目录占位符
doc.add_heading('目录', level=1)
toc_paragraphs = [
    '1. 简介',
    '2. 总体描述',
    '3. 系统架构',
    '4. 组件设计',
    '5. 数据设计',
    '6. 接口设计',
    '7. 安全设计',
    '8. 非功能需求',
    '9. 部署架构',
    '10. 附录'
]

for item in toc_paragraphs:
    p = doc.add_paragraph(item, style='List Bullet')
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_page_break()

# 第1章: 简介
doc.add_heading('1. 简介', level=1)

doc.add_heading('1.1 目的', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('本文档描述 EmoHealer 情绪疗愈平台的高层设计,包括系统架构、组件设计、数据模型和技术规范。HLD 文档旨在为开发团队、架构师和技术决策者提供系统的技术设计指导。'), '微软雅黑', 11)

doc.add_heading('1.2 范围', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('EmoHealer 是一个基于 AIGC(人工智能生成内容)技术的情绪疗愈平台,为 18-35 岁用户提供 7×24 小时专业且温暖的情绪支持服务。系统包括以下核心组件:'), '微软雅黑', 11)

modules = [
    '• 用户认证与管理模块',
    '• AI 智能对话系统(基于 CBT 认知行为疗法)',
    '• 多模态情绪识别(文本、语音、表情)',
    '• 个性化疗愈方案生成',
    '• 情绪日记与数据分析',
    '• 危机预警与干预机制',
    '• 心理测评与咨询师预约',
    '• 系统管理后台'
]

for module in modules:
    p = doc.add_paragraph(module)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_heading('1.3 定义、缩写和术语', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('关键技术术语定义:'), '微软雅黑', 11, False)

term_table = doc.add_table(rows=11, cols=2)
term_table.style = 'Table Grid'

terms = [
    ['HLD', 'High-Level Design (高层设计)'],
    ['SRS', 'Software Requirements Specification (软件需求规格说明书)'],
    ['CBT', 'Cognitive Behavioral Therapy (认知行为疗法)'],
    ['LLM', 'Large Language Model (大语言模型)'],
    ['RAG', 'Retrieval-Augmented Generation (检索增强生成)'],
    ['MVP', 'Minimum Viable Product (最小可行产品)'],
    ['API', 'Application Programming Interface (应用程序接口)'],
    ['REST', 'Representational State Transfer (表述性状态转移)'],
    ['ORM', 'Object-Relational Mapping (对象关系映射)'],
    ['SQLAlchemy', 'Python ORM 框架'],
    ['Pydantic', 'Python 数据验证库']
]

for i, (abbr, full) in enumerate(terms):
    row = term_table.rows[i]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(abbr), '微软雅黑', 10, True)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(full), '微软雅黑', 10, False)

doc.add_heading('1.4 参考文献', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('主要参考文献:'), '微软雅黑', 11, False)

refs = [
    '1. EmoHealer 软件需求规格说明书 (SRS)',
    '2. FastAPI 官方文档: https://fastapi.tiangolo.com/',
    '3. SQLAlchemy 文档: https://docs.sqlalchemy.org/',
    '4. WebSocket 协议规范: RFC 6455',
    '5. MySQL 8.0 参考手册',
    '6. Docker 官方文档',
    '7. 心理健康服务相关法律法规'
]

for ref in refs:
    p = doc.add_paragraph(ref)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_page_break()

# 第2章: 总体描述
doc.add_heading('2. 总体描述', level=1)

doc.add_heading('2.1 产品视角', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('EmoHealer 采用前后端分离的三层架构:'), '微软雅黑', 11)

layers = [
    '1. 表现层 (Presentation Layer):',
    '   - Web 浏览器客户端',
    '   - 响应式 HTML5/JavaScript 界面',
    '   - ECharts 数据可视化组件',
    '   - WebSocket 实时通信客户端',
    '',
    '2. 应用层 (Application Layer):',
    '   - FastAPI RESTful API 服务',
    '   - WebSocket 实时通信服务',
    '   - AI 智能体服务(多角色 LLM 集成)',
    '   - 业务逻辑处理服务',
    '   - 认证与授权服务',
    '',
    '3. 数据层 (Data Layer):',
    '   - MySQL 8.0 关系型数据库',
    '   - SQLAlchemy ORM 数据访问层',
    '   - 数据缓存策略(可选 Redis)'
]

for layer in layers:
    p = doc.add_paragraph(layer)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_heading('2.2 产品功能', level=2)
function_table = doc.add_table(rows=10, cols=3)
function_table.style = 'Table Grid'

# 表头
header_row = function_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('功能模块'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('功能描述'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('优先级'), '微软雅黑', 10, True)

functions = [
    ['用户认证', '注册、登录、登出、Token 管理、密码修改', '高'],
    ['AI 对话', '智能对话、情绪识别、危机检测、多角色回复', '高'],
    ['情绪分析', '文本/语音/表情情绪识别、情绪趋势分析、情绪报告', '高'],
    ['疗愈方案', 'AI 生成个性化方案、任务执行、完成率统计', '中'],
    ['情绪日记', '创建、查询、更新、删除日记(软删除)', '中'],
    ['危机预警', '关键词检测、预警级别评估、管理员处理', '高'],
    ['心理测评', 'PHQ-9/GAD-7 测评提交、历史查询', '中'],
    ['咨询预约', '提交预约、预约列表、状态管理', '中'],
    ['系统管理', '数据驾驶舱、用户管理、Token 统计、操作日志', '中']
]

for i, func in enumerate(functions):
    row = function_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(func[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(func[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(func[2]), '微软雅黑', 10, False)

doc.add_heading('2.3 用户特征', level=2)
user_table = doc.add_table(rows=4, cols=4)
user_table.style = 'Table Grid'

# 表头
header_row = user_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('用户类型'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('特征描述'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('需求'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[3].paragraphs[0].add_run('使用场景'), '微软雅黑', 10, True)

users = [
    ['普通用户', '18-35岁,面临工作压力、情感困扰、情绪管理需求', '完整功能访问', '随时访问、移动友好、隐私保护'],
    ['管理员', '负责系统运维、数据监控、危机处理', '管理权限', '数据看板、危机预警处理'],
    ['咨询师', '接受预约,提供专业心理咨询', '预约管理', '查看预约、提供咨询']
]

for i, user in enumerate(users):
    row = user_table.rows[i+1]
    for j, cell_text in enumerate(user):
        set_chinese_font(row.cells[j].paragraphs[0].add_run(cell_text), '微软雅黑', 10, False)

doc.add_page_break()

# 第3章: 系统架构
doc.add_heading('3. 系统架构', level=1)

doc.add_heading('3.1 架构概览', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('EmoHealer 采用经典的三层架构(Tiered Architecture),结合微服务理念:'), '微软雅黑', 11)

architecture_desc = [
    '',
    '1. 表现层: 负责用户交互和界面展示',
    '   - Web 客户端: 通过浏览器访问 http://localhost:8092/emohealer',
    '   - 管理后台: 通过 /admin 路径访问',
    '   - 使用 WebSocket 实现实时双向通信',
    '',
    '2. 应用层: 负责业务逻辑处理和 API 服务',
    '   - API Gateway: 统一入口,处理 CORS、认证',
    '   - 认证服务: 用户注册/登录/Token 验证',
    '   - 业务服务: 对话、情绪分析、疗愈方案等',
    '   - AI 服务: 集成 LLM,多角色智能体',
    '   - WebSocket 管理器: 实时消息推送',
    '',
    '3. 数据层: 负责数据持久化和查询',
    '   - MySQL 8.0: 主数据存储',
    '   - 可选 Redis: 缓存热点数据',
    '',
    '4. 外部服务: LLM API 集成',
    '   - 文心一言: 百度提供的 LLM 服务',
    '   - ChatGLM: 智谱 AI 提供的 LLM 服务',
    '   - OpenAI: 备选 LLM 服务'
]

for desc in architecture_desc:
    p = doc.add_paragraph(desc)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_heading('3.2 系统设计原则', level=2)
principle_table = doc.add_table(rows=10, cols=3)
principle_table.style = 'Table Grid'

# 表头
header_row = principle_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('设计原则'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('说明'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('应用场景'), '微软雅黑', 10, True)

principles = [
    ['关注点分离', '表现层、应用层、数据层职责明确', '整体架构设计'],
    ['单一职责', '每个组件只负责一个功能领域', '服务模块划分'],
    ['开闭原则', '对扩展开放,对修改关闭', 'AI 服务可替换'],
    ['依赖倒置', '依赖抽象而非具体实现', 'ORM 抽象层'],
    ['RESTful 设计', '使用标准 HTTP 方法和状态码', 'API 设计'],
    ['安全优先', '认证、加密、输入验证贯穿全层', '安全设计'],
    ['可扩展性', '支持横向扩展,增加服务器实例', '部署架构'],
    ['高可用性', '无单点故障,支持故障恢复', '部署和运维'],
    ['用户数据隔离', '严格的数据权限控制', '数据访问层']
]

for i, principle in enumerate(principles):
    row = principle_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(principle[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(principle[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(principle[2]), '微软雅黑', 10, False)

doc.add_page_break()

# 第4章: 组件设计
doc.add_heading('4. 组件设计', level=1)

doc.add_heading('4.1 前端组件', level=2)
frontend_table = doc.add_table(rows=9, cols=3)
frontend_table.style = 'Table Grid'

# 表头
header_row = frontend_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('组件名称'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('功能描述'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('技术实现'), '微软雅黑', 10, True)

frontend_components = [
    ['页面路由器', '单页应用(SPA)路由切换', 'JavaScript showPage() 函数'],
    ['认证管理器', 'Token 存储、验证、自动刷新', 'localStorage + 拦截器'],
    ['聊天界面', '消息显示、输入框、情绪标签选择', 'HTML + CSS + JS'],
    ['WebSocket 客户端', '实时通信、心跳保活', 'WebSocket API'],
    ['情绪图表', 'ECharts 情绪趋势图', 'ECharts line/bar 图表'],
    ['疗愈方案', '方案展示、任务执行、进度跟踪', 'HTML + JS 状态管理'],
    ['情绪日记', '日记列表、编辑器、删除功能', '表单 + 列表组件'],
    ['心理测评', 'PHQ-9/GAD-7 量表、结果展示', '表单 + 结果页面'],
    ['危机预警', '预警弹窗、求助热线显示', '模态框组件']
]

for i, component in enumerate(frontend_components):
    row = frontend_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(component[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(component[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(component[2]), '微软雅黑', 10, False)

doc.add_heading('4.2 后端组件', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('API 路由结构:'), '微软雅黑', 11)

route_table = doc.add_table(rows=9, cols=3)
route_table.style = 'Table Grid'

# 表头
header_row = route_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('路由模块'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('路径前缀'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('主要端点'), '微软雅黑', 10, True)

routes = [
    ['API 路由', '/api', '所有业务 API'],
    ['认证路由', '/api/auth', '登录/注册/登出/验证'],
    ['聊天路由', '/api/chat', '发送消息/历史记录'],
    ['情绪路由', '/api/emotion', '情绪分析/趋势/报告'],
    ['疗愈方案路由', '/api/plans', '生成/查询方案'],
    ['日记路由', '/api/diary', 'CRUD 操作'],
    ['测评路由', '/api/assessment', '提交/查询测评'],
    ['预约路由', '/api/appointment', '提交/查询预约'],
    ['管理路由', '/api/admin', '驾驶舱/用户管理/统计']
]

for i, route in enumerate(routes):
    row = route_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(route[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(route[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(route[2]), '微软雅黑', 10, False)

doc.add_heading('4.3 数据库组件', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('数据库表清单:'), '微软雅黑', 11)

db_table = doc.add_table(rows=12, cols=4)
db_table.style = 'Table Grid'

# 表头
header_row = db_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('表名'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('用途'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('记录量预估'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[3].paragraphs[0].add_run('索引策略'), '微软雅黑', 10, True)

db_tables = [
    ['user', '用户基本信息', '10,000+', 'PRIMARY, UNIQUE(username)'],
    ['user_token', '用户登录 Token', '50,000+', 'INDEX(user_id), INDEX(token)'],
    ['chat_record', '对话记录', '1,000,000+', 'INDEX(user_id), INDEX(emotion_type), INDEX(created_at)'],
    ['emotion_log', '情绪日志', '2,000,000+', 'INDEX(user_id), INDEX(emotion_type), INDEX(created_at)'],
    ['healing_plan', '疗愈方案', '100,000+', 'INDEX(user_id), INDEX(plan_date)'],
    ['crisis_alert', '危机预警记录', '1,000+', 'INDEX(user_id), INDEX(alert_level), INDEX(is_handled)'],
    ['psychological_assessment', '心理测评', '50,000+', 'INDEX(user_id), INDEX(assessment_type)'],
    ['consultation_appointment', '咨询预约', '10,000+', 'INDEX(user_id), INDEX(appointment_time), INDEX(status)'],
    ['emotion_diary', '情绪日记', '200,000+', 'INDEX(user_id), INDEX(created_at), INDEX(emotion_type)'],
    ['system_config', '系统配置', '50', 'PRIMARY(config_key)'],
    ['token_usage', 'Token 消耗记录', '500,000+', 'INDEX(user_id), INDEX(model_provider), INDEX(created_at)'],
    ['operation_log', '操作日志', '1,000,000+', 'INDEX(user_id), INDEX(operation_type), INDEX(created_at)']
]

for i, table_info in enumerate(db_tables):
    row = db_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(table_info[0]), '微软雅黑', 9, True)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(table_info[1]), '微软雅黑', 9, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(table_info[2]), '微软雅黑', 9, False)
    set_chinese_font(row.cells[3].paragraphs[0].add_run(table_info[3]), '微软雅黑', 9, False)

doc.add_page_break()

# 第5章: 数据设计
doc.add_heading('5. 数据设计', level=1)

doc.add_heading('5.1 数据模型', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('核心数据实体:'), '微软雅黑', 11)

entity_table = doc.add_table(rows=11, cols=4)
entity_table.style = 'Table Grid'

# 表头
header_row = entity_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('实体'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('主要属性'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('关系'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[3].paragraphs[0].add_run('业务规则'), '微软雅黑', 10, True)

entities = [
    ['User', 'id, username, password_hash, nickname, avatar_url, email, phone', '1:N user_token, chat_record, emotion_log, healing_plan, crisis_alert, assessment, appointment, diary', '用户名唯一,密码SHA256加密'],
    ['UserToken', 'user_id, token, expires_at', 'N:1 User', 'Token有效期7天,过期自动失效'],
    ['ChatRecord', 'user_id, user_message, ai_reply, emotion_type, emotion_score, is_crisis', 'N:1 User, 1:1 CrisisAlert', '保存完整对话历史'],
    ['EmotionLog', 'user_id, emotion_type, emotion_score, confidence, source, context_tags', 'N:1 User', '支持多源情绪识别'],
    ['HealingPlan', 'user_id, plan_date, tasks(JSON), completion_rate, status, ai_summary', 'N:1 User', '每日一个方案,任务JSON格式'],
    ['CrisisAlert', 'user_id, alert_level, reason, keywords(JSON), is_handled', 'N:1 User, N:1 ChatRecord', '必须管理员处理'],
    ['PsychologicalAssessment', 'user_id, assessment_type, total_score, level, answers(JSON)', 'N:1 User', 'PHQ-9/GAD-7标准化量表'],
    ['ConsultationAppointment', 'user_id, appointment_time, status, consultation_type', 'N:1 User', '预约时间唯一'],
    ['EmotionDiary', 'user_id, title, content, mood_tags(JSON), emotion_type, is_archived', 'N:1 User', '软删除机制'],
    ['SystemConfig', 'config_key, config_value, description', '-', '全局配置'],
    ['TokenUsage', 'user_id, model_provider, prompt_tokens, completion_tokens, cost', 'N:1 User', '统计AI成本']
]

for i, entity in enumerate(entities):
    row = entity_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(entity[0]), '微软雅黑', 9, True)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(entity[1]), '微软雅黑', 9, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(entity[2]), '微软雅黑', 9, False)
    set_chinese_font(row.cells[3].paragraphs[0].add_run(entity[3]), '微软雅黑', 9, False)

doc.add_heading('5.2 数据库架构', level=2)

storage_table = doc.add_table(rows=3, cols=2)
storage_table.style = 'Table Grid'

storage_data = [
    ['存储引擎', 'InnoDB: 所有表使用 InnoDB 引擎\n- 支持事务(ACID)\n- 行级锁定\n- 支持外键约束\n- 崩溃恢复能力强'],
    ['字符集', 'utf8mb4: 支持完整的 Unicode 字符(包括 emoji), 向后兼容 utf8'],
    ['排序规则', 'utf8mb4_unicode_ci: 不区分大小写, Unicode 排序规则']
]

for i, (key, value) in enumerate(storage_data):
    row = storage_table.rows[i]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(key), '微软雅黑', 10, True)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(value), '微软雅黑', 10, False)
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(4.5)

doc.add_page_break()

# 第6章: 接口设计
doc.add_heading('6. 接口设计', level=1)

doc.add_heading('6.1 用户接口', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('前端页面列表:'), '微软雅黑', 11)

page_table = doc.add_table(rows=10, cols=4)
page_table.style = 'Table Grid'

# 表头
header_row = page_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('页面'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('路由路径'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('主要功能'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[3].paragraphs[0].add_run('认证要求'), '微软雅黑', 10, True)

pages = [
    ['主页(聊天页面)', '#chat', 'AI对话、情绪选择、消息历史', '必须登录'],
    ['登录页', '#login', '用户名/密码登录', '否'],
    ['注册页', '#register', '新用户注册', '否'],
    ['用户中心', '#profile', '个人信息、统计数据', '必须登录'],
    ['情绪报告', '#emotion-report', '情绪趋势图表、分析报告', '必须登录'],
    ['疗愈方案', '#healing-plan', '今日方案、任务执行', '必须登录'],
    ['情绪日记', '#diary', '日记列表、编辑、删除', '必须登录'],
    ['心理测评', '#assessment', 'PHQ-9/GAD-7量表', '必须登录'],
    ['咨询师预约', '#appointment', '预约表单、预约列表', '必须登录'],
    ['关于我们', '#about', '系统介绍', '否']
]

for i, page in enumerate(pages):
    row = page_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(page[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(page[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(page[2]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[3].paragraphs[0].add_run(page[3]), '微软雅黑', 10, False)

doc.add_heading('6.2 API 接口', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('RESTful API 设计原则:'), '微软雅黑', 11)

api_principles = [
    '• 资源导向: URL 表示资源,HTTP 方法表示操作',
    '• 统一响应格式: 所有接口返回统一 JSON 结构',
    '• 状态码使用: 遵循 HTTP 标准状态码',
    '• 版本控制: 通过 URL 前缀 /api 控制版本'
]

for principle in api_principles:
    p = doc.add_paragraph(principle)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

p = doc.add_paragraph()
set_chinese_font(p.add_run('统一响应格式:'), '微软雅黑', 11, True)

# 成功响应示例
p = doc.add_paragraph()
code_run = p.add_run('成功响应:')
set_chinese_font(code_run, '微软雅黑', 11, True)

code_block = doc.add_paragraph()
code_block.paragraph_format.left_indent = Cm(0.5)
code_run = code_block.add_run('{')
set_chinese_font(code_run, 'Consolas', 9, False)

code_block = doc.add_paragraph()
code_block.paragraph_format.left_indent = Cm(0.5)
code_run = code_block.add_run('  "code": 200,')
set_chinese_font(code_run, 'Consolas', 9, False)

code_block = doc.add_paragraph()
code_block.paragraph_format.left_indent = Cm(0.5)
code_run = code_block.add_run('  "message": "操作成功",')
set_chinese_font(code_run, 'Consolas', 9, False)

code_block = doc.add_paragraph()
code_block.paragraph_format.left_indent = Cm(0.5)
code_run = code_block.add_run('  "data": { ... }')
set_chinese_font(code_run, 'Consolas', 9, False)

code_block = doc.add_paragraph()
code_block.paragraph_format.left_indent = Cm(0.5)
code_run = code_block.add_run('}')
set_chinese_font(code_run, 'Consolas', 9, False)

doc.add_page_break()

# 第7章: 安全设计
doc.add_heading('7. 安全设计', level=1)

doc.add_heading('7.1 认证和授权', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('Token 认证机制:'), '微软雅黑', 11, True)

token_mechanisms = [
    '1. Token 生成:',
    '   - 用户登录成功后生成随机 Token',
    '   - Token 长度: 256 位随机字符串',
    '   - 存储到 user_token 表',
    '   - 设置过期时间: 当前时间 + 7 天',
    '',
    '2. Token 验证:',
    '   - 所有需要认证的接口通过 Authorization 头传递 Token',
    '   - 中间件验证 Token 是否存在且未过期',
    '   - 验证通过后从 user_token 表获取 user_id',
    '',
    '3. Token 失效:',
    '   - 用户登出时删除 Token',
    '   - Token 过期自动失效(通过 expires_at 字段判断)',
    '   - 修改密码时使所有旧 Token 失效'
]

for mechanism in token_mechanisms:
    p = doc.add_paragraph(mechanism)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_heading('7.2 密码安全', level=2)
password_table = doc.add_table(rows=5, cols=2)
password_table.style = 'Table Grid'

password_data = [
    ['加密算法', 'SHA-256'],
    ['盐值', '固定盐值 "EmoHealer2026"'],
    ['加密公式', 'hash = SHA256(password + salt)'],
    ['存储', '只存储哈希值,不存储明文密码'],
    ['传输加密', 'HTTPS (生产环境)']
]

for i, (key, value) in enumerate(password_data):
    row = password_table.rows[i]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(key), '微软雅黑', 10, True)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(value), '微软雅黑', 10, False)

doc.add_heading('7.3 输入验证', level=2)

validation_table = doc.add_table(rows=5, cols=3)
validation_table.style = 'Table Grid'

# 表头
header_row = validation_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('验证项'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('规则'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('实现'), '微软雅黑', 10, True)

validations = [
    ['用户名', '3-50字符,字母数字下划线', 'HTML5 pattern'],
    ['密码', '6-50字符', 'HTML5 pattern'],
    ['邮箱', '标准邮箱格式', 'HTML5 type="email"'],
    ['手机号', '11位数字', 'HTML5 pattern'],
    ['日记内容', '1-10000字符', 'HTML5 maxlength']
]

for i, validation in enumerate(validations):
    row = validation_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(validation[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(validation[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(validation[2]), '微软雅黑', 10, False)

doc.add_page_break()

# 第8章: 非功能需求
doc.add_heading('8. 非功能需求', level=1)

doc.add_heading('8.1 性能需求', level=2)
performance_table = doc.add_table(rows=6, cols=3)
performance_table.style = 'Table Grid'

# 表头
header_row = performance_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('指标'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('目标'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('测量方法'), '微软雅黑', 10, True)

performances = [
    ['页面加载时间', '< 3秒', '浏览器开发者工具'],
    ['API响应时间(数据库查询)', '< 500ms', 'API日志时间戳差'],
    ['AI对话响应时间', '< 5秒', '前端-后端时间差'],
    ['WebSocket连接建立', '< 2秒', '连接事件时间'],
    ['情绪分析响应', '< 1秒', 'API日志'],
    ['数据库查询响应', '< 100ms(简单查询)', 'MySQL慢查询日志']
]

for i, perf in enumerate(performances):
    row = performance_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(perf[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(perf[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(perf[2]), '微软雅黑', 10, False)

doc.add_heading('8.2 可扩展性', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('水平扩展:'), '微软雅黑', 11, True)

h_scale = [
    '• 应用服务器: 支持 Nginx 负载均衡,多个 FastAPI 实例,无状态设计可任意扩展',
    '• 数据库: MySQL 主从复制,读写分离(写主库,读从库),分库分表(用户维度)'
]

for item in h_scale:
    p = doc.add_paragraph(item)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_heading('8.3 可靠性', level=2)
reliability_table = doc.add_table(rows=4, cols=3)
reliability_table.style = 'Table Grid'

# 表头
header_row = reliability_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('故障场景'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('容错措施'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('恢复策略'), '微软雅黑', 10, True)

reliabilities = [
    ['数据库宕机', '连接池重试', '自动重连 + 错误提示'],
    ['LLM API 超时', '规则模板回退', '使用预设回复'],
    ['网络中断', 'WebSocket重连', '指数退避重连'],
    ['服务器崩溃', '日志持久化', '重启后日志恢复']
]

for i, reliability in enumerate(reliabilities):
    row = reliability_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(reliability[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(reliability[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(reliability[2]), '微软雅黑', 10, False)

doc.add_heading('8.4 可用性', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('系统可用性目标:'), '微软雅黑', 11, True)

availability = [
    '• 目标: ≥ 99.5%',
    '• 年停机时间: < 43.8 小时'
]

for item in availability:
    p = doc.add_paragraph(item)
    if p.runs:
        set_chinese_font(p.runs[0], '微软雅黑', 11, False)

doc.add_page_break()

# 第9章: 部署架构
doc.add_heading('9. 部署架构', level=1)

doc.add_heading('9.1 部署环境', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('开发环境:'), '微软雅黑', 11, True)

dev_env = [
    '部署方式: 本地运行',
    '',
    '# 启动后端',
    'cd backend',
    'python main.py  # 默认端口: 8092',
    '',
    '# 启动前端',
    'cd frontend',
    'python -m http.server 5000',
    '',
    '# 访问',
    '# 前端: http://localhost:5000/emohealer2.html',
    '# 后端API: http://localhost:8092',
    '# API文档: http://localhost:8092/docs'
]

for env in dev_env:
    p = doc.add_paragraph(env)
    if p.runs:
        set_chinese_font(p.runs[0], 'Consolas', 10, False)

p = doc.add_paragraph()
set_chinese_font(p.add_run('数据库: 本地 MySQL 8.0'), '微软雅黑', 11, False)

doc.add_heading('9.2 基础设施要求', level=2)
infra_table = doc.add_table(rows=6, cols=2)
infra_table.style = 'Table Grid'

# 表头
header_row = infra_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('软件'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('版本/用途'), '微软雅黑', 10, True)

infra = [
    ['Python', '3.10+ / 后端运行时'],
    ['MySQL', '8.0+ / 数据库'],
    ['Nginx', '1.20+ / 反向代理/负载均衡'],
    ['Docker', '20.10+ / 容器化(可选)'],
    ["Let's Encrypt", '- / SSL 证书']
]

for i, item in enumerate(infra):
    row = infra_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(item[0]), '微软雅黑', 10, True)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(item[1]), '微软雅黑', 10, False)

doc.add_page_break()

# 第10章: 附录
doc.add_heading('10. 附录', level=1)

doc.add_heading('10.1 配置参数', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('后端配置:'), '微软雅黑', 11, True)

config_table = doc.add_table(rows=7, cols=3)
config_table.style = 'Table Grid'

# 表头
header_row = config_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('配置项'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('环境变量'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('默认值/说明'), '微软雅黑', 10, True)

configs = [
    ['数据库主机', 'DB_HOST', 'localhost / MySQL服务器地址'],
    ['数据库端口', 'DB_PORT', '3306 / MySQL端口'],
    ['数据库用户', 'DB_USER', 'root / MySQL用户名'],
    ['数据库密码', 'DB_PASSWORD', '19891213 / MySQL密码'],
    ['数据库名', 'DB_NAME', 'emohealer / 数据库名称'],
    ['服务主机', 'HOST', '0.0.0.0 / 监听地址'],
    ['服务端口', 'PORT', '8092 / FastAPI端口']
]

for i, config in enumerate(configs):
    row = config_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(config[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(config[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(config[2]), '微软雅黑', 10, False)

doc.add_heading('10.2 错误代码', level=2)
error_table = doc.add_table(rows=8, cols=4)
error_table.style = 'Table Grid'

# 表头
header_row = error_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('错误代码'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('HTTP状态'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('描述'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[3].paragraphs[0].add_run('处理建议'), '微软雅黑', 10, True)

errors = [
    ['200', '200', '成功', '-'],
    ['400', '400', '请求参数错误', '检查参数格式'],
    ['401', '401', '未授权/Token无效', '重新登录'],
    ['403', '403', '账号禁用', '联系管理员'],
    ['404', '404', '资源不存在', '检查资源ID'],
    ['422', '422', '数据验证失败', '检查输入数据'],
    ['500', '500', '服务器内部错误', '联系技术支持'],
    ['503', '503', '服务不可用', '稍后重试']
]

for i, error in enumerate(errors):
    row = error_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(error[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(error[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(error[2]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[3].paragraphs[0].add_run(error[3]), '微软雅黑', 10, False)

# 文档版本历史
doc.add_page_break()
doc.add_heading('文档版本历史', level=1)

history_table = doc.add_table(rows=2, cols=4)
history_table.style = 'Table Grid'

# 表头
header_row = history_table.rows[0]
set_chinese_font(header_row.cells[0].paragraphs[0].add_run('版本'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[1].paragraphs[0].add_run('日期'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[2].paragraphs[0].add_run('修订人'), '微软雅黑', 10, True)
set_chinese_font(header_row.cells[3].paragraphs[0].add_run('修订内容'), '微软雅黑', 10, True)

history = [
    ['v1.0', '2026-03-19', 'Claude', '初始版本,基于EmoHealer项目代码分析生成']
]

for i, item in enumerate(history):
    row = history_table.rows[i+1]
    set_chinese_font(row.cells[0].paragraphs[0].add_run(item[0]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[1].paragraphs[0].add_run(item[1]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[2].paragraphs[0].add_run(item[2]), '微软雅黑', 10, False)
    set_chinese_font(row.cells[3].paragraphs[0].add_run(item[3]), '微软雅黑', 10, False)

# 文档结束
doc.add_paragraph()
p = doc.add_paragraph('**文档结束**')
set_chinese_font(p.runs[0], '微软雅黑', 12, True)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
output_file = 'c:/Users/18746/Desktop/biyeshixi/EmoHealer_HLD_v1.0.docx'
doc.save(output_file)

print(f'✅ 文档已生成: {output_file}')
print(f'📄 文档包含以下章节:')
print(f'   1. 简介')
print(f'   2. 总体描述')
print(f'   3. 系统架构')
print(f'   4. 组件设计')
print(f'   5. 数据设计')
print(f'   6. 接口设计')
print(f'   7. 安全设计')
print(f'   8. 非功能需求')
print(f'   9. 部署架构')
print(f'   10. 附录')
