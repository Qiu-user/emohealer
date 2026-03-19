#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成 EmoHealer HLD 文档 - 简化版
基于 EmoHealer_HLD_v1.0.md 生成 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# 设置页面
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)

# 设置中文字体函数
def add_paragraph_with_font(doc, text, font_name='微软雅黑', size=12, bold=False):
    """添加段落并设置字体"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

# 添加标题
title = doc.add_heading('EmoHealer 情绪疗愈平台', level=0)
run = title.runs[0]
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(20)

subtitle = doc.add_heading('高层设计文档 (HLD)', level=1)
run = subtitle.runs[0]
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(16)

# 添加文档信息
add_paragraph_with_font(doc, '项目名称: EmoHealer AI 情绪疗愈平台', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '文档版本: v1.0', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '创建日期: 2026-03-19', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '作者: 基于 EmoHealer 项目代码分析生成', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '对应 SRS: EmoHealer_功能需求文档.md', '微软雅黑', 11, False)

doc.add_page_break()

# 添加目录
add_paragraph_with_font(doc, '目录', '微软雅黑', 16, True)
add_paragraph_with_font(doc, '1. 简介', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '2. 总体描述', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '3. 系统架构', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '4. 组件设计', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '5. 数据设计', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '6. 接口设计', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '7. 安全设计', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '8. 非功能需求', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '9. 部署架构', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '10. 附录', '微软雅黑', 11, False)

doc.add_page_break()

# 第1章: 简介
add_paragraph_with_font(doc, '1. 简介', '微软雅黑', 14, True)
add_paragraph_with_font(doc, '1.1 目的', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '本文档描述 EmoHealer 情绪疗愈平台的高层设计,包括系统架构、组件设计、数据模型和技术规范。HLD 文档旨在为开发团队、架构师和技术决策者提供系统的技术设计指导。', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '1.2 范围', '微软雅黑', 12, True)
add_paragraph_with_font(doc, 'EmoHealer 是一个基于 AIGC(人工智能生成内容)技术的情绪疗愈平台,为 18-35 岁用户提供 7×24 小时专业且温暖的情绪支持服务。', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '系统包括以下核心组件:', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 用户认证与管理模块', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• AI 智能对话系统(基于 CBT 认知行为疗法)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 多模态情绪识别(文本、语音、表情)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 个性化疗愈方案生成', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪日记与数据分析', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 危机预警与干预机制', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 心理测评与咨询师预约', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 系统管理后台', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '1.3 定义、缩写和术语', '微软雅黑', 12, True)
add_paragraph_with_font(doc, 'HLD: High-Level Design (高层设计)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'SRS: Software Requirements Specification (软件需求规格说明书)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'CBT: Cognitive Behavioral Therapy (认知行为疗法)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'LLM: Large Language Model (大语言模型)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'API: Application Programming Interface (应用程序接口)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'REST: Representational State Transfer (表述性状态转移)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'ORM: Object-Relational Mapping (对象关系映射)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'SQLAlchemy: Python ORM 框架', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'Pydantic: Python 数据验证库', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '1.4 参考文献', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '1. EmoHealer 软件需求规格说明书 (SRS)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '2. FastAPI 官方文档: https://fastapi.tiangolo.com/', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '3. SQLAlchemy 文档: https://docs.sqlalchemy.org/', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '4. WebSocket 协议规范: RFC 6455', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '5. MySQL 8.0 参考手册', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '6. Docker 官方文档', '微软雅黑', 11, False)

doc.add_page_break()

# 第2章: 总体描述
add_paragraph_with_font(doc, '2. 总体描述', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '2.1 产品视角', '微软雅黑', 12, True)
add_paragraph_with_font(doc, 'EmoHealer 采用前后端分离的三层架构:', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '1. 表现层 (Presentation Layer):', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - Web 浏览器客户端', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 响应式 HTML5/JavaScript 界面', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - ECharts 数据可视化组件', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - WebSocket 实时通信客户端', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '2. 应用层 (Application Layer):', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - FastAPI RESTful API 服务', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - WebSocket 实时通信服务', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - AI 智能体服务(多角色 LLM 集成)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 业务逻辑处理服务', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 认证与授权服务', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '3. 数据层 (Data Layer):', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - MySQL 8.0 关系型数据库', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - SQLAlchemy ORM 数据访问层', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 数据缓存策略(可选 Redis)', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '2.2 产品功能', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '系统提供以下核心功能:', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 用户认证: 注册、登录、登出、Token 管理、密码修改 (优先级: 高)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• AI 对话: 智能对话、情绪识别、危机检测、多角色回复 (优先级: 高)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪分析: 文本/语音/表情情绪识别、情绪趋势分析、情绪报告 (优先级: 高)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 疗愈方案: AI 生成个性化方案、任务执行、完成率统计 (优先级: 中)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪日记: 创建、查询、更新、删除日记(软删除) (优先级: 中)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 危机预警: 关键词检测、预警级别评估、管理员处理 (优先级: 高)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 心理测评: PHQ-9/GAD-7 测评提交、历史查询 (优先级: 中)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 咨询预约: 提交预约、预约列表、状态管理 (优先级: 中)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 系统管理: 数据驾驶舱、用户管理、Token 统计、操作日志 (优先级: 中)', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '2.3 用户特征', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '普通用户: 18-35岁,面临工作压力、情感困扰、情绪管理需求', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '管理员: 负责系统运维、数据监控、危机处理', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '咨询师: 接受预约,提供专业心理咨询', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '2.4 约束条件', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '技术约束:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 浏览器兼容性: 需支持 Chrome 80+、Firefox 75+、Edge 80+', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据库: 必须使用 MySQL 8.0+,支持 InnoDB 引擎', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• Python 版本: 后端需要 Python 3.10+', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• API 版本: RESTful API 遵循 HTTP/1.1 规范', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• WebSocket: 必须支持 RFC 6455 协议', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• LLM 响应时间: AI 对话响应时间 < 5 秒', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '业务约束:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• Token 有效期: 访问 Token 有效期为 7 天', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 每日对话限制: 每日最多 100 次对话', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 密码加密: 必须使用 SHA256 加密存储(带盐值)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据隔离: 用户数据严格隔离,无法访问其他用户数据', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 危机响应: 检测到危机必须在 1 秒内推送预警', '微软雅黑', 11, False)

doc.add_page_break()

# 第3章: 系统架构
add_paragraph_with_font(doc, '3. 系统架构', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '3.1 架构概览', '微软雅黑', 12, True)
add_paragraph_with_font(doc, 'EmoHealer 采用经典的三层架构(Tiered Architecture),结合微服务理念:', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '1. 表现层: 负责用户交互和界面展示', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - Web 客户端: 通过浏览器访问 http://localhost:8092/emohealer', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 管理后台: 通过 /admin 路径访问', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 使用 WebSocket 实现实时双向通信', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '2. 应用层: 负责业务逻辑处理和 API 服务', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - API Gateway: 统一入口,处理 CORS、认证', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 认证服务: 用户注册/登录/Token 验证', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 业务服务: 对话、情绪分析、疗愈方案等', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - AI 服务: 集成 LLM,多角色智能体', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - WebSocket 管理器: 实时消息推送', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '3. 数据层: 负责数据持久化和查询', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - MySQL 8.0: 主数据存储', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 可选 Redis: 缓存热点数据', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '4. 外部服务: LLM API 集成', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - 文心一言: 百度提供的 LLM 服务', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - ChatGLM: 智谱 AI 提供的 LLM 服务', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '   - OpenAI: 备选 LLM 服务', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '3.2 系统设计原则', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '• 关注点分离: 表现层、应用层、数据层职责明确', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 单一职责: 每个组件只负责一个功能领域', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 开闭原则: 对扩展开放,对修改关闭', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 依赖倒置: 依赖抽象而非具体实现', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• RESTful 设计: 使用标准 HTTP 方法和状态码', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 安全优先: 认证、加密、输入验证贯穿全层', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 可扩展性: 支持横向扩展,增加服务器实例', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 高可用性: 无单点故障,支持故障恢复', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 用户数据隔离: 严格的数据权限控制', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '3.3 技术栈', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '表现层:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• HTML/CSS/JavaScript: 标准前端技术', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• ECharts 5.4+: 数据可视化图表库', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• Web Speech API: 浏览器语音识别 API', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• WebSocket RFC 6455: 实时通信协议', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '应用层:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• Python 3.10+: 后端编程语言', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• FastAPI 0.104.0+: Web 框架', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• uvicorn 0.23.0+: ASGI 服务器', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• SQLAlchemy 2.0+: ORM 框架', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• PyMySQL 1.1.0+: MySQL 数据库驱动', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• Pydantic 2.0+: 数据验证库', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• python-jose: JWT Token 处理', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '数据层:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• MySQL 8.0+: 关系型数据库', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• InnoDB: 存储引擎', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• Redis 6.0+: 缓存数据库(可选)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '外部服务:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 文心一言: 百度 LLM API', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• ChatGLM: 智谱 AI LLM API', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• OpenAI: OpenAI GPT API', '微软雅黑', 11, False)

doc.add_page_break()

# 第4章: 组件设计
add_paragraph_with_font(doc, '4. 组件设计', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '4.1 前端组件', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '• 页面路由器: 单页应用(SPA)路由切换 - JavaScript showPage() 函数', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 认证管理器: Token 存储、验证、自动刷新 - localStorage + 拦截器', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 聊天界面: 消息显示、输入框、情绪标签选择 - HTML + CSS + JS', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• WebSocket 客户端: 实时通信、心跳保活 - WebSocket API', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪图表: ECharts 情绪趋势图 - ECharts line/bar 图表', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 疗愈方案: 方案展示、任务执行、进度跟踪 - HTML + JS 状态管理', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪日记: 日记列表、编辑器、删除功能 - 表单 + 列表组件', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 心理测评: PHQ-9/GAD-7 量表、结果展示 - 表单 + 结果页面', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 危机预警: 预警弹窗、求助热线显示 - 模态框组件', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '4.2 后端组件', '微软雅黑', 12, True)
add_paragraph_with_font(doc, 'API 路由结构:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• /api: 所有业务 API', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/auth: 登录/注册/登出/验证', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/chat: 发送消息/历史记录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/emotion: 情绪分析/趋势/报告', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/plans: 生成/查询方案', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/diary: CRUD 操作', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/assessment: 提交/查询测评', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/appointment: 提交/查询预约', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /api/admin: 驾驶舱/用户管理/统计', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• /ws: 实时聊天连接', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '4.3 数据库组件', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '数据库表清单:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• user: 用户基本信息 (10,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• user_token: 用户登录 Token (50,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• chat_record: 对话记录 (1,000,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• emotion_log: 情绪日志 (2,000,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• healing_plan: 疗愈方案 (100,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• crisis_alert: 危机预警记录 (1,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• psychological_assessment: 心理测评 (50,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• consultation_appointment: 咨询预约 (10,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• emotion_diary: 情绪日记 (200,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• system_config: 系统配置 (50 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• token_usage: Token 消耗记录 (500,000+ 记录)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• operation_log: 操作日志 (1,000,000+ 记录)', '微软雅黑', 11, False)

doc.add_page_break()

# 第5章: 数据设计
add_paragraph_with_font(doc, '5. 数据设计', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '5.1 数据模型', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '核心数据实体:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• User: id, username, password_hash, nickname, avatar_url, email, phone - 用户名唯一,密码SHA256加密', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• UserToken: user_id, token, expires_at - Token有效期7天,过期自动失效', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• ChatRecord: user_id, user_message, ai_reply, emotion_type, emotion_score, is_crisis - 保存完整对话历史', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• EmotionLog: user_id, emotion_type, emotion_score, confidence, source, context_tags - 支持多源情绪识别', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• HealingPlan: user_id, plan_date, tasks(JSON), completion_rate, status, ai_summary - 每日一个方案,任务JSON格式', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• CrisisAlert: user_id, alert_level, reason, keywords(JSON), is_handled - 必须管理员处理', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• PsychologicalAssessment: user_id, assessment_type, total_score, level, answers(JSON) - PHQ-9/GAD-7标准化量表', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• ConsultationAppointment: user_id, appointment_time, status, consultation_type - 预约时间唯一', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• EmotionDiary: user_id, title, content, mood_tags(JSON), emotion_type, is_archived - 软删除机制', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• SystemConfig: config_key, config_value, description - 全局配置', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• TokenUsage: user_id, model_provider, prompt_tokens, completion_tokens, cost - 统计AI成本', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• OperationLog: user_id, operation_type, resource_type, resource_id, ip_address - 审计追踪', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '5.2 数据库架构', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '存储引擎: InnoDB - 支持事务(ACID)、行级锁定、支持外键约束、崩溃恢复能力强', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '字符集: utf8mb4 - 支持完整的 Unicode 字符(包括 emoji), 向后兼容 utf8', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '排序规则: utf8mb4_unicode_ci - 不区分大小写, Unicode 排序规则', '微软雅黑', 11, False)

doc.add_page_break()

# 第6章: 接口设计
add_paragraph_with_font(doc, '6. 接口设计', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '6.1 用户接口', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '前端页面列表:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 主页(聊天页面) #chat: AI对话、情绪选择、消息历史 - 必须登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 登录页 #login: 用户名/密码登录 - 否', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 注册页 #register: 新用户注册 - 否', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 用户中心 #profile: 个人信息、统计数据 - 必须登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪报告 #emotion-report: 情绪趋势图表、分析报告 - 必须登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 疗愈方案 #healing-plan: 今日方案、任务执行 - 必须登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪日记 #diary: 日记列表、编辑、删除 - 必须登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 心理测评 #assessment: PHQ-9/GAD-7量表 - 必须登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 咨询师预约 #appointment: 预约表单、预约列表 - 必须登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 关于我们 #about: 系统介绍 - 否', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '6.2 API 接口', '微软雅黑', 12, True)
add_paragraph_with_font(doc, 'RESTful API 设计原则:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 资源导向: URL 表示资源,HTTP 方法表示操作', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 统一响应格式: 所有接口返回统一 JSON 结构', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 状态码使用: 遵循 HTTP 标准状态码', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 版本控制: 通过 URL 前缀 /api 控制版本', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '统一响应格式:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '成功响应:', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '{', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '  "code": 200,', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '  "message": "操作成功",', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '  "data": { ... }', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '}', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '错误响应:', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '{', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '  "code": 400,', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '  "message": "请求参数错误",', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '  "data": null', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '}', '微软雅黑', 10, False)

add_paragraph_with_font(doc, '主要 API 端点:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '认证接口:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• POST /api/auth/login - 用户登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• POST /api/auth/register - 用户注册', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• POST /api/auth/logout - 用户登出', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• GET /api/auth/verify - 验证Token', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '聊天接口:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• POST /api/chat/send - 发送消息', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• GET /api/chat/history - 获取对话历史', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• GET /api/chat/history/{user_id} - 获取用户对话历史', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '情绪接口:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• POST /api/emotion/analyze - 分析情绪', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• GET /api/emotion/trend - 获取情绪趋势', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• GET /api/emotion/report/{user_id} - 获取情绪报告', '微软雅黑', 11, False)
add_paragraph_with_font(doc, 'WebSocket 接口:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• ws://localhost:8092/ws/chat?token=xxx - 实时聊天', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• ws://localhost:8092/ws/admin - 管理后台', '微软雅黑', 11, False)

doc.add_page_break()

# 第7章: 安全设计
add_paragraph_with_font(doc, '7. 安全设计', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '7.1 认证和授权', '微软雅黑', 12, True)
add_paragraph_with_font(doc, 'Token 认证机制:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '1. Token 生成: 用户登录成功后生成随机 Token,Token 长度256位随机字符串,存储到user_token表,设置过期时间为当前时间+7天', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '2. Token 验证: 所有需要认证的接口通过Authorization头传递Token,中间件验证Token是否存在且未过期,验证通过后从user_token表获取user_id', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '3. Token 失效: 用户登出时删除Token,Token过期自动失效(通过expires_at字段判断),修改密码时使所有旧Token失效', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '7.2 密码安全', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '• 加密算法: SHA-256', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 盐值: 固定盐值 "EmoHealer2026"', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 加密公式: hash = SHA256(password + salt)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 存储: 只存储哈希值,不存储明文密码', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 传输加密: HTTPS (生产环境)', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '7.3 输入验证', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '前端验证:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 用户名: 3-50字符,字母数字下划线 - HTML5 pattern', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 密码: 6-50字符 - HTML5 pattern', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 邮箱: 标准邮箱格式 - HTML5 type="email"', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 手机号: 11位数字 - HTML5 pattern', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 日记内容: 1-10000字符 - HTML5 maxlength', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '后端验证: 使用 Pydantic 模型进行数据验证,参数化查询防SQL注入,最小权限原则', '微软雅黑', 11, False)

doc.add_page_break()

# 第8章: 非功能需求
add_paragraph_with_font(doc, '8. 非功能需求', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '8.1 性能需求', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '• 页面加载时间: < 3秒 - 浏览器开发者工具', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• API响应时间(数据库查询): < 500ms - API日志时间戳差', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• AI对话响应时间: < 5秒 - 前端-后端时间差', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• WebSocket连接建立: < 2秒 - 连接事件时间', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 情绪分析响应: < 1秒 - API日志', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据库查询响应: < 100ms(简单查询) - MySQL慢查询日志', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '8.2 可扩展性', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '水平扩展:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 应用服务器: 支持 Nginx 负载均衡,多个 FastAPI 实例,无状态设计可任意扩展', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据库: MySQL 主从复制,读写分离(写主库,读从库),分库分表(用户维度)', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '8.3 可靠性', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '• 数据库宕机: 连接池重试 - 自动重连 + 错误提示', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• LLM API 超时: 规则模板回退 - 使用预设回复', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 网络中断: WebSocket重连 - 指数退避重连', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 服务器崩溃: 日志持久化 - 重启后日志恢复', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '8.4 可用性', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '• 目标: ≥ 99.5%', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 年停机时间: < 43.8 小时', '微软雅黑', 11, False)

doc.add_page_break()

# 第9章: 部署架构
add_paragraph_with_font(doc, '9. 部署架构', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '9.1 部署环境', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '开发环境:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '部署方式: 本地运行', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '# 启动后端', '微软雅黑', 10, False)
add_paragraph_with_font(doc, 'cd backend', '微软雅黑', 10, False)
add_paragraph_with_font(doc, 'python main.py  # 默认端口: 8092', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '# 启动前端', '微软雅黑', 10, False)
add_paragraph_with_font(doc, 'cd frontend', '微软雅黑', 10, False)
add_paragraph_with_font(doc, 'python -m http.server 5000', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '# 访问', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '# 前端: http://localhost:5000/emohealer2.html', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '# 后端API: http://localhost:8092', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '# API文档: http://localhost:8092/docs', '微软雅黑', 10, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '数据库: 本地 MySQL 8.0', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '9.2 基础设施要求', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '软件依赖:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• Python 3.10+ - 后端运行时', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• MySQL 8.0+ - 数据库', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• Nginx 1.20+ - 反向代理/负载均衡', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• Docker 20.10+ - 容器化(可选)', '微软雅黑', 11, False)
add_paragraph_with_font(doc, "• Let's Encrypt - SSL 证书", '微软雅黑', 11, False)
add_paragraph_with_font(doc, '', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '网络配置:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 80 - HTTP - TCP', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 443 - HTTPS - TCP', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 8092 - FastAPI - TCP', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 3306 - MySQL - TCP(内网)', '微软雅黑', 11, False)

doc.add_page_break()

# 第10章: 附录
add_paragraph_with_font(doc, '10. 附录', '微软雅黑', 14, True)

add_paragraph_with_font(doc, '10.1 配置参数', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '后端配置:', '微软雅黑', 11, True)
add_paragraph_with_font(doc, '• 数据库主机 (DB_HOST): localhost - MySQL服务器地址', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据库端口 (DB_PORT): 3306 - MySQL端口', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据库用户 (DB_USER): root - MySQL用户名', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据库密码 (DB_PASSWORD): 19891213 - MySQL密码', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 数据库名 (DB_NAME): emohealer - 数据库名称', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 服务主机 (HOST): 0.0.0.0 - 监听地址', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 服务端口 (PORT): 8092 - FastAPI端口', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '10.2 错误代码', '微软雅黑', 12, True)
add_paragraph_with_font(doc, '• 200 - 成功', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 400 - 请求参数错误 - 检查参数格式', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 401 - 未授权/Token无效 - 重新登录', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 403 - 账号禁用 - 联系管理员', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 404 - 资源不存在 - 检查资源ID', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 422 - 数据验证失败 - 检查输入数据', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 500 - 服务器内部错误 - 联系技术支持', '微软雅黑', 11, False)
add_paragraph_with_font(doc, '• 503 - 服务不可用 - 稍后重试', '微软雅黑', 11, False)

add_paragraph_with_font(doc, '10.3 部署检查清单', '微软雅黑', 12, True)
checklist = [
    '• 服务器配置满足最低要求',
    '• Python 3.10+ 已安装',
    '• MySQL 8.0+ 已安装并初始化',
    '• 所有依赖包已安装(requirements.txt)',
    '• 环境变量已配置(.env文件)',
    '• 数据库连接测试通过',
    '• API服务启动成功(python main.py)',
    '• Nginx反向代理已配置',
    '• SSL证书已安装',
    '• 域名DNS解析已配置',
    '• 防火墙规则已设置',
    '• API文档可访问(/docs)',
    '• 前端页面可访问(/emohealer)',
    '• WebSocket连接测试通过',
    '• 用户注册/登录功能测试通过',
    '• AI对话功能测试通过',
    '• 情绪分析功能测试通过',
    '• 危机预警功能测试通过',
    '• 监控系统已配置',
    '• 备份策略已设置',
    '• 日志系统已启用'
]

for item in checklist:
    add_paragraph_with_font(doc, item, '微软雅黑', 11, False)

# 文档版本历史
doc.add_page_break()
add_paragraph_with_font(doc, '文档版本历史', '微软雅黑', 14, True)
add_paragraph_with_font(doc, '• v1.0 - 2026-03-19 - Claude - 初始版本,基于EmoHealer项目代码分析生成', '微软雅黑', 11, False)

# 文档结束
doc.add_paragraph()
add_paragraph_with_font(doc, '**文档结束**', '微软雅黑', 12, True)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
output_file = 'c:/Users/18746/Desktop/biyeshixi/EmoHealer_HLD_v1.0.docx'
doc.save(output_file)

print(f'[OK] 文档已生成: {output_file}')
print(f'[DOC] 文档包含以下章节:')
print(f'   1. 简介')
print(f'   2. 总体描述')
print(f'   3. 系统架构')
print(f'   4. 组件设计')
print(f'   5. 数据设计')
print(f'   6. 接口设计')
print(f'   7. 安全设计')
print(f'   8. 非功能需求')
print(f'   9. 部署架构')
print(f'   10. 附录')
