from docx import Document

doc = Document('EmoHealer_HLD_v1.0.docx')
print(f'段落数: {len(doc.paragraphs)}')
print(f'章节数: {len([p for p in doc.paragraphs if p.style.name.startswith("Heading")])}')
print(f'表格数: {len(doc.tables)}')

print('\n主要章节:')
for i, p in enumerate(doc.paragraphs[:50]):
    if p.style.name.startswith('Heading'):
        print(f'  {p.style.name}: {p.text[:50]}')
